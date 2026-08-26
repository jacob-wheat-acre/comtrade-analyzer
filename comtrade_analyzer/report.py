"""
report.py — Word (.docx) report generator for COMTRADE relay event analysis.

Structure mirrors pq_report.py from the PQ analyzer:
  generate_report()        → compile EventRecord + analysis into a summary dict
  generate_word_report()   → assemble the .docx document
  Private _word_xxx()      → one function per document section

Requires:  pip install python-docx
Optional:  plots embedded from save_path PNG files produced by plotting.py
"""

from __future__ import annotations

import logging
import os
import tempfile
from pathlib import Path
from typing import Dict, List, Optional

import numpy as np

from .data_model import EventRecord
from .analysis import (
    compute_event_summary,
    detect_fault_inception,
    detect_trip_time,
    detect_digital_transitions,
    estimate_dc_offset,
    compute_sequence_components,
)

log = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# python-docx — optional import guard (same pattern as pq_report.py)
# ---------------------------------------------------------------------------

try:
    from docx import Document as _DocxDocument
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

# ---------------------------------------------------------------------------
# Color palette — neutral / utility-agnostic
# ---------------------------------------------------------------------------

def _rgb(r, g, b):
    return RGBColor(r, g, b) if _DOCX_AVAILABLE else None

_NAVY     = _rgb(0x1A, 0x3A, 0x6B)   # header / section titles
_MID_BLUE = _rgb(0x2C, 0x6E, 0xA6)   # secondary headings
_TRIP_RED = _rgb(0xC0, 0x39, 0x2B)   # trip / fault markers
_PASS_CLR = _rgb(0x1E, 0x7A, 0x1E)   # operated / detected (green)
_WARN_CLR = _rgb(0xE8, 0x77, 0x22)   # caution (orange)
_GRAY_CLR = _rgb(0xF2, 0xF2, 0xF2)   # unused / N/A


# ---------------------------------------------------------------------------
# Low-level docx helpers (identical API to pq_report.py)
# ---------------------------------------------------------------------------

def _cell_shade(cell, hex_color: str) -> None:
    """Apply background fill to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_col_widths(table, widths_cm: list) -> None:
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)


def _bold(para, text: str, color=None, size_pt: int = 11):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def _normal(para, text: str, color=None, size_pt: int = 11):
    run = para.add_run(text)
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def _section_heading(doc, title: str) -> None:
    p = doc.add_paragraph()
    _bold(p, title, color=_NAVY, size_pt=11)


def _body(doc, text: str, size_pt: int = 10) -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(size_pt)


# ---------------------------------------------------------------------------
# 1.  generate_report() — build the summary dict
# ---------------------------------------------------------------------------

def generate_report(record: EventRecord) -> dict:
    """
    Compile EventRecord analysis into a structured summary dict.

    This is the single source of truth consumed by both generate_word_report()
    and the stdout print_report() in main.py.
    """
    summary = compute_event_summary(record)

    # Fault inception details
    fault_idx = detect_fault_inception(record)

    # Trip details
    trip_result = detect_trip_time(record)
    trip_idx, trip_ch = trip_result if trip_result else (None, None)

    # Digital channel operations
    digital_ops = []
    for name, data in record.digital_channels.items():
        trans = detect_digital_transitions(record, name)
        for idx in trans["rising"]:
            t_s = float(record.time[idx])
            rel_ms = (t_s - record.trigger_time) * 1000
            digital_ops.append({
                "channel": name,
                "event":   "ASSERT (0→1)",
                "time_s":  t_s,
                "rel_trigger_ms": rel_ms,
                "sample":  idx,
            })
        for idx in trans["falling"]:
            t_s = float(record.time[idx])
            rel_ms = (t_s - record.trigger_time) * 1000
            digital_ops.append({
                "channel": name,
                "event":   "RELEASE (1→0)",
                "time_s":  t_s,
                "rel_trigger_ms": rel_ms,
                "sample":  idx,
            })
    digital_ops.sort(key=lambda x: x["time_s"])

    # DC offset on the faulted phase
    dc_offset_info: dict = {}
    if fault_idx is not None:
        spc = record.samples_per_cycle()
        for name, data in record.analog_channels.items():
            units = record.analog_info[name].units.upper()
            if not any(u in units for u in ("A", "AMP")):
                continue
            post = data[fault_idx:]
            if len(post) >= spc:
                dc, tau = estimate_dc_offset(post, spc)
                if abs(dc) > 0.05 * float(np.max(np.abs(post[:spc]))):
                    dc_offset_info[name] = {"dc_a": dc, "tau_s": tau}

    # Sequence components snapshot at trigger
    def _get(*names):
        for n in names:
            ch = record.get_channel(n)
            if ch is not None:
                return ch
        return None

    ia = _get("IA", "Ia", "I_A")
    ib = _get("IB", "Ib", "I_B")
    ic = _get("IC", "Ic", "I_C")

    seq_at_fault: dict = {}
    if ia is not None and ib is not None and ic is not None and fault_idx is not None:
        spc = record.samples_per_cycle()
        end = min(fault_idx + spc * 2, len(record.time))
        mid = (fault_idx + end) // 2
        if mid >= spc:
            i0, i1, i2 = compute_sequence_components(ia, ib, ic, spc)
            seq_at_fault = {
                "I1_pos": float(np.nanmean(i1[fault_idx:end])),
                "I2_neg": float(np.nanmean(i2[fault_idx:end])),
                "I0_zero": float(np.nanmean(i0[fault_idx:end])),
            }

    return {
        "event_summary":   summary,
        "digital_ops":     digital_ops,
        "dc_offset":       dc_offset_info,
        "sequence":        seq_at_fault,
        "fault_index":     fault_idx,
        "trip_index":      trip_idx,
        "trip_channel":    trip_ch,
        "record_metadata": record.metadata,
    }


# ---------------------------------------------------------------------------
# 2.  Word document section helpers
# ---------------------------------------------------------------------------

def _word_event_info_table(doc, record: EventRecord, report: dict,
                           location: str = "", device_type: str = "") -> None:
    """Site / event metadata table — mirrors _word_site_info_table in pq_report.py."""
    meta  = record.metadata
    start = meta.get("start_time")
    trig  = meta.get("trigger_time_abs")
    start_str = start.strftime("%Y-%m-%d  %H:%M:%S.%f")[:-3] if start else "—"
    trig_str  = trig.strftime("%Y-%m-%d  %H:%M:%S.%f")[:-3]  if trig  else "—"

    rows = [
        ("Station",            meta.get("station_name", "—")),
        ("Device / Relay ID",  meta.get("rec_dev_id",   "—")),
    ]
    if location:
        rows.append(("Location / Substation", location))
    if device_type:
        rows.append(("Relay / Device Type", device_type))
    rows += [
        ("Recording Start",    start_str),
        ("Trigger Time",       trig_str),
        ("Duration",           f"{record.duration_s()*1000:.1f} ms"),
        ("Sample Rate",        f"{record.sample_rate:.0f} Hz  "
                               f"({record.samples_per_cycle()} samples/cycle)"),
        ("Channels",           f"{len(record.analog_channels)} analog  /  "
                               f"{len(record.digital_channels)} digital"),
        ("COMTRADE Revision",  meta.get("rev_year", "—")),
        ("Data Format",        meta.get("file_type", "—")),
    ]

    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [5.5, 11.0])
    for i, (label, value) in enumerate(rows):
        lc, vc = tbl.rows[i].cells
        _cell_shade(lc, "E8F1FA")
        lc.paragraphs[0].add_run(label).bold = True
        vc.paragraphs[0].add_run(str(value))
    doc.add_paragraph()


def _word_operations_table(doc, record: EventRecord, report: dict) -> None:
    """
    Event operations summary — mirrors the compliance table in pq_report.py.

    Columns: Parameter | Measured / Detail | Status
    """
    summary   = report["event_summary"]
    fault_t   = summary.get("fault_inception_s")
    trip_t    = summary.get("trip_time_s")
    trig_t    = record.trigger_time
    fault_dur = summary.get("fault_duration_s")
    delay_ms  = summary.get("trip_delay_ms")
    ftype     = summary.get("fault_type", "UNKNOWN")
    dig_ops   = report.get("digital_ops", [])

    _fault_labels = {
        "SLG":  "Single Line-to-Ground (SLG)",
        "LL":   "Line-to-Line (LL)",
        "LLG":  "Double Line-to-Ground (LLG)",
        "3PH":  "Three-Phase (3PH)",
        "UNKNOWN": "Undetermined",
    }

    hdr_p = doc.add_paragraph()
    _bold(hdr_p, "Event Operations Summary", color=_NAVY, size_pt=12)

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [5.5, 9.0, 2.5])

    # Header row
    for cell, text in zip(tbl.rows[0].cells, ["Parameter", "Measured / Detail", "Status"]):
        _cell_shade(cell, "1A3A6B")
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)

    def _add_row(param: str, measured: str, status: str, status_color=None,
                 row_shade: str = None):
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(param).font.size = Pt(10)
        row.cells[1].paragraphs[0].add_run(measured).font.size = Pt(10)
        sr = row.cells[2].paragraphs[0].add_run(status)
        sr.bold = True
        sr.font.size = Pt(10)
        if status_color:
            sr.font.color.rgb = status_color
        if row_shade:
            for c in row.cells:
                _cell_shade(c, row_shade)

    # Fault type
    _add_row(
        "Fault Type",
        _fault_labels.get(ftype, ftype),
        "Classified" if ftype != "UNKNOWN" else "N/A",
        _PASS_CLR if ftype != "UNKNOWN" else _GRAY_CLR,
    )

    # Fault inception
    if fault_t is not None:
        rel = (fault_t - trig_t) * 1000
        meas = f"{fault_t*1000:.2f} ms  ({rel:+.1f} ms from trigger)"
        _add_row("Fault Inception", meas, "Detected", _PASS_CLR)
    else:
        _add_row("Fault Inception", "Not detected in current window", "N/A", _GRAY_CLR)

    # Protection elements operated
    operated = [op["channel"] for op in dig_ops if op["event"].startswith("ASSERT")]
    if operated:
        op_str = ", ".join(dict.fromkeys(operated))  # deduplicate, preserve order
        _add_row("Elements Operated", op_str, "Operated", _PASS_CLR)
    else:
        _add_row("Elements Operated", "No digital assertions detected", "None", _GRAY_CLR)

    # Trip time
    if trip_t is not None:
        rel = (trip_t - trig_t) * 1000
        ch  = report.get("trip_channel", "")
        meas = f"{trip_t*1000:.2f} ms  ({rel:+.1f} ms from trigger)"
        if ch:
            meas += f"  [{ch}]"
        _add_row("Trip Time", meas, "Tripped", _PASS_CLR)
    else:
        _add_row("Trip Time", "No TRIP signal detected", "N/A", _GRAY_CLR)

    # Fault clearing time
    if fault_dur is not None:
        _add_row("Fault Clearing Time",
                 f"{fault_dur*1000:.1f} ms  (inception → trip)",
                 f"{fault_dur*1000:.0f} ms",
                 _PASS_CLR if fault_dur < 0.200 else _WARN_CLR,
                 "FFF8E8" if fault_dur >= 0.200 else None)
    else:
        _add_row("Fault Clearing Time", "Cannot compute (inception or trip missing)", "N/A",
                 _GRAY_CLR)

    # Relay operate time / trip delay
    if delay_ms is not None:
        _add_row("Relay Operate Time (Trip Delay)",
                 f"{delay_ms:.1f} ms  (IEEE time-delay grading reference: varies by coordination)",
                 f"{delay_ms:.0f} ms",
                 _PASS_CLR if delay_ms < 100 else _WARN_CLR)
    else:
        _add_row("Relay Operate Time", "N/A", "N/A", _GRAY_CLR)

    doc.add_paragraph()


def _word_peak_quantities(doc, record: EventRecord, report: dict) -> None:
    """Table of peak analog channel values."""
    summary = report["event_summary"]
    currents = summary.get("max_currents", {})
    voltages = summary.get("max_voltages", {})

    _section_heading(doc, "Peak Measured Quantities")

    if not currents and not voltages:
        _body(doc, "No analog channel data available.")
        return

    rows_data: list = []
    for name, peak in currents.items():
        info  = record.analog_info.get(name)
        units = info.units if info else "A"
        rows_data.append((name, f"{peak:.2f}", units, "Current"))
    for name, peak in voltages.items():
        info  = record.analog_info.get(name)
        units = info.units if info else "V"
        rows_data.append((name, f"{peak:.2f}", units, "Voltage"))

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [4.0, 4.5, 3.0, 5.0])
    for cell, text in zip(tbl.rows[0].cells, ["Channel", "Peak Magnitude", "Units", "Type"]):
        _cell_shade(cell, "E8F1FA")
        r = cell.paragraphs[0].add_run(text)
        r.bold = True
        r.font.size = Pt(10)

    for i, (ch, val, units, ctype) in enumerate(rows_data):
        row = tbl.add_row()
        if i % 2 == 1:
            for c in row.cells:
                _cell_shade(c, "F7FBFF")
        row.cells[0].paragraphs[0].add_run(ch).font.size = Pt(10)
        row.cells[1].paragraphs[0].add_run(val).font.size = Pt(10)
        row.cells[2].paragraphs[0].add_run(units).font.size = Pt(10)
        row.cells[3].paragraphs[0].add_run(ctype).font.size = Pt(10)

    doc.add_paragraph()


def _word_waveform(doc, plot_path: Optional[str]) -> None:
    """Embed the waveform PNG produced by plotting.py."""
    _section_heading(doc, "Waveform Record")

    if plot_path and os.path.isfile(plot_path):
        doc.add_picture(plot_path, width=Cm(16))
        p = doc.add_paragraph()
        r = p.add_run(
            "Figure: Three-phase currents, voltages, and digital status.  "
            "Red dashed line = fault inception.  Green dotted line = trip.  "
            "Orange line = trigger."
        )
        r.font.size = Pt(8)
        r.italic = True
    else:
        _body(doc,
            "Waveform plot not available. Run with Save Plots enabled to embed the figure.",
            size_pt=9)

    doc.add_paragraph()


def _word_triage(doc, triage: dict) -> None:
    """Render a colored priority banner immediately after the title block."""
    priority = triage.get("priority", 3)
    labels   = triage.get("labels", [])
    notes    = triage.get("notes", [])

    # Banner background and text colors by priority
    if priority == 1:
        bg_hex, hdr_color, hdr_text = "C0392B", _rgb(0xFF, 0xFF, 0xFF), "PRIORITY 1 — IMMEDIATE REVIEW REQUIRED"
    elif priority == 2:
        bg_hex, hdr_color, hdr_text = "E87722", _rgb(0xFF, 0xFF, 0xFF), "PRIORITY 2 — ROUTINE REVIEW"
    else:
        bg_hex, hdr_color, hdr_text = "1E7A1E", _rgb(0xFF, 0xFF, 0xFF), "PRIORITY 3 — ARCHIVE (no review required)"

    # Single-row banner table
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_col_widths(tbl, [16.0])
    cell = tbl.cell(0, 0)
    _cell_shade(cell, bg_hex)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(hdr_text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = hdr_color

    # Flag detail rows (one row per flag)
    if labels:
        detail_tbl = doc.add_table(rows=len(labels), cols=2)
        detail_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        _set_col_widths(detail_tbl, [4.0, 12.0])
        for i, (lbl, note) in enumerate(zip(labels, notes)):
            row = detail_tbl.rows[i]
            _cell_shade(row.cells[0], "F2F2F2")
            lp = row.cells[0].paragraphs[0]
            lr = lp.add_run(lbl)
            lr.bold = True
            lr.font.size = Pt(9)
            np_ = row.cells[1].paragraphs[0]
            nr = np_.add_run(note)
            nr.font.size = Pt(9)

    doc.add_paragraph()


def _word_phasor(doc, phasor_png: Optional[str]) -> None:
    """Embed the phasor diagram PNG produced by plotting.plot_phasors()."""
    _section_heading(doc, "Phasor Diagram")

    if phasor_png and os.path.isfile(phasor_png):
        doc.add_picture(phasor_png, width=Cm(16))
        p = doc.add_paragraph()
        r = p.add_run(
            "Figure: Voltage phasors (Va, Vb, Vc), current phasors (Ia, Ib, Ic), and "
            "symmetrical sequence currents (I₁, I₂, I₀).  "
            "Faded arrows = pre-fault.  Bold arrows = fault window (one cycle after inception).  "
            "Reference: Va (fault) = 0°."
        )
        r.font.size = Pt(8)
        r.italic = True
    else:
        _body(doc,
            "Phasor diagram not available. Enable Phasor Plot in the options to embed the figure.",
            size_pt=9)

    doc.add_paragraph()


def _word_digital_log(doc, record: EventRecord, report: dict) -> None:
    """Chronological table of all digital channel state changes."""
    _section_heading(doc, "Digital Operations Log")

    dig_ops = report.get("digital_ops", [])
    if not dig_ops:
        _body(doc, "No digital channel transitions detected.")
        doc.add_paragraph()
        return

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [4.0, 4.5, 4.5, 3.5])
    for cell, text in zip(tbl.rows[0].cells,
                          ["Channel", "Event", "Time (ms)", "Rel. Trigger (ms)"]):
        _cell_shade(cell, "E8F1FA")
        r = cell.paragraphs[0].add_run(text)
        r.bold = True
        r.font.size = Pt(10)

    for i, op in enumerate(dig_ops):
        row = tbl.add_row()
        if i % 2 == 1:
            for c in row.cells:
                _cell_shade(c, "F7FBFF")
        is_trip = ("TRIP" in op["channel"].upper() or op["channel"].upper().startswith("TR")) \
                  and op["event"].startswith("ASSERT")

        name_run = row.cells[0].paragraphs[0].add_run(op["channel"])
        name_run.font.size = Pt(10)
        if is_trip:
            name_run.bold = True
            name_run.font.color.rgb = _TRIP_RED

        row.cells[1].paragraphs[0].add_run(op["event"]).font.size = Pt(10)
        row.cells[2].paragraphs[0].add_run(f"{op['time_s']*1000:.2f}").font.size = Pt(10)
        rel = op["rel_trigger_ms"]
        row.cells[3].paragraphs[0].add_run(f"{rel:+.2f}").font.size = Pt(10)

    doc.add_paragraph()


def _word_analysis_notes(doc, record: EventRecord, report: dict) -> None:
    """DC offset, sequence components, fault type rationale."""
    _section_heading(doc, "Analysis Notes")

    fault_type = report["event_summary"].get("fault_type", "UNKNOWN")
    type_labels = {
        "SLG": "Single Line-to-Ground (A-phase to ground is the most common type, "
               "accounting for approximately 70–80% of all transmission faults).",
        "LL":  "Line-to-Line. Two phases are faulted together without significant ground involvement. "
               "Common cause: phase-to-phase contact during high-wind or galloping conductor events.",
        "LLG": "Double Line-to-Ground. Two phases faulted to ground. "
               "Often initiated by a single conductor contacting ground and subsequently involving "
               "an adjacent phase.",
        "3PH": "Three-Phase (balanced). All three phases involved. Least common type; "
               "often caused by equipment failure or improper switching. Produces no negative "
               "or zero-sequence current.",
        "UNKNOWN": "Fault type could not be classified from available channel data.",
    }
    _body(doc, f"Fault classification:  {type_labels.get(fault_type, fault_type)}")

    # DC offset
    dc_info = report.get("dc_offset", {})
    if dc_info:
        doc.add_paragraph()
        p = doc.add_paragraph()
        _bold(p, "DC Offset:", size_pt=10)
        for ch, vals in dc_info.items():
            dc_a = vals["dc_a"]
            tau  = vals["tau_s"]
            tau_str = f"{tau*1000:.1f} ms" if tau > 0 else "not estimated"
            _body(doc,
                f"  {ch}: DC component ≈ {dc_a:.1f} A  |  time constant τ ≈ {tau_str}. "
                "DC offset is caused by the point-on-wave of fault inception relative to the "
                "voltage zero crossing. Full offset (50% of peak current added as DC) "
                "is worst-case for breaker interruption duty.")

    # Sequence components
    seq = report.get("sequence", {})
    if seq:
        doc.add_paragraph()
        p = doc.add_paragraph()
        _bold(p, "Symmetrical Sequence Components (post-fault, one-cycle average):", size_pt=10)
        _body(doc,
            f"  I₁ (Positive) = {seq.get('I1_pos', 0):.1f} A  |  "
            f"I₂ (Negative) = {seq.get('I2_neg', 0):.1f} A  |  "
            f"I₀ (Zero) = {seq.get('I0_zero', 0):.1f} A")
        i1 = seq.get("I1_pos", 0)
        i2 = seq.get("I2_neg", 0)
        i0 = seq.get("I0_zero", 0)
        if i1 > 0:
            neg_ratio = i2 / i1 * 100 if i1 > 0 else 0
            zer_ratio = i0 / i1 * 100 if i1 > 0 else 0
            _body(doc,
                f"  Negative-to-positive ratio: {neg_ratio:.0f}%  "
                f"(< 5% expected for balanced 3-phase faults; "
                f"significant I₂ indicates phase-to-phase or ground fault involvement). "
                f"Zero-to-positive ratio: {zer_ratio:.0f}%  "
                f"(significant I₀ confirms ground fault path).")

    doc.add_paragraph()


def _word_reclose_operations(doc, feeder: dict) -> None:
    """
    Recloser shot-by-shot table and fault location estimate.

    Only rendered when feeder analysis data is present in the report dict.
    """
    from .feeder_analysis import RecloserSequence, RecloserShot

    seq: RecloserSequence = feeder.get("reclose_sequence")
    if seq is None or seq.total_shots == 0:
        return

    _section_heading(doc, "Recloser / Feeder Relay Operations")

    # Outcome sentence
    outcome_color = _PASS_CLR if "SUCCESSFUL" in seq.final_outcome else (
        _TRIP_RED if "LOCKED OUT" in seq.final_outcome else _WARN_CLR
    )
    p = doc.add_paragraph()
    _bold(p, f"Outcome:  {seq.final_outcome}", color=outcome_color, size_pt=10)

    permanence_text = {
        "TEMPORARY":     "Fault is TEMPORARY — cleared on reclose.  "
                         "Likely caused by tree contact, animal, or transient overvoltage.",
        "PERMANENT":     "Fault is PERMANENT — device locked out.  "
                         "A line patrol or feeder sectionalizing is required to locate and clear the fault.",
        "INDETERMINATE": "Fault permanence could not be determined from this record.  "
                         "Check subsequent SCADA / relay logs for reclose outcome.",
    }
    _body(doc, permanence_text.get(seq.fault_type, ""), size_pt=10)
    doc.add_paragraph()

    # Shot-by-shot table
    shot_hdr = doc.add_paragraph()
    _bold(shot_hdr, "Operation Sequence", size_pt=10)

    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [1.5, 3.0, 3.0, 4.0, 4.0, 3.5])
    for cell, text in zip(tbl.rows[0].cells,
                          ["Shot", "Element", "Type", "Operate Time (ms)",
                           "Dead Time (ms)", "Outcome"]):
        _cell_shade(cell, "1A3A6B")
        r = cell.paragraphs[0].add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)

    for shot in seq.shots:
        row = tbl.add_row()
        is_lockout = shot.outcome == "LOCKOUT"
        is_success = shot.successful and not shot.post_reclose_fault

        if is_lockout:
            for c in row.cells:
                _cell_shade(c, "FFE8E8")
        elif is_success:
            for c in row.cells:
                _cell_shade(c, "E8F4E8")

        row.cells[0].paragraphs[0].add_run(str(shot.shot_number)).font.size = Pt(10)
        row.cells[1].paragraphs[0].add_run(shot.element or "—").font.size = Pt(10)

        type_run = row.cells[2].paragraphs[0].add_run(shot.shot_type)
        type_run.font.size = Pt(10)
        if shot.shot_type == "FAST":
            type_run.font.color.rgb = _PASS_CLR

        op_ms = f"{shot.operate_time_ms:.1f}" if shot.operate_time_ms is not None else "N/A"
        row.cells[3].paragraphs[0].add_run(op_ms).font.size = Pt(10)

        dead = f"{shot.dead_time_ms:.0f}" if shot.dead_time_ms is not None else "N/A"
        row.cells[4].paragraphs[0].add_run(dead).font.size = Pt(10)

        outcome_run = row.cells[5].paragraphs[0].add_run(shot.outcome)
        outcome_run.font.size = Pt(10)
        outcome_run.bold = True
        if is_lockout:
            outcome_run.font.color.rgb = _TRIP_RED
        elif is_success:
            outcome_run.font.color.rgb = _PASS_CLR

    doc.add_paragraph()

    # Fault location estimate
    loc = feeder.get("fault_location")
    if loc and loc.get("estimated_miles") is not None:
        loc_hdr = doc.add_paragraph()
        _bold(loc_hdr, "Fault Location Estimate", size_pt=10)
        _body(doc,
            f"Faulted phase: {loc['faulted_phase']}.  "
            f"Estimated distance: ~{loc['estimated_miles']:.1f} miles from the substation.  "
            f"Fault current: {loc['i_fault_rms_a']:.0f} A RMS.  "
            f"Apparent fault impedance: {loc['z_fault_ohm']:.2f} Ω at "
            f"{loc['feeder_z_ohm_per_mile']:.3f} Ω/mile.  "
            f"Accuracy: {loc['accuracy_note']}.",
            size_pt=10)
        doc.add_paragraph()

    # HIF screen
    hif = feeder.get("hif_screen", {})
    if hif.get("hif_suspect"):
        hif_p = doc.add_paragraph()
        _bold(hif_p, "High-Impedance Fault (HIF) Screen — SUSPECT", color=_WARN_CLR, size_pt=10)
        _body(doc,
            f"Current increase (ΔI = {hif['delta_current_a']:.1f} A) is below the HIF threshold "
            f"of {hif['threshold_a']:.1f} A.  {hif['note']}",
            size_pt=10)
        doc.add_paragraph()

    # Voltage recovery per shot
    recoveries = feeder.get("voltage_recoveries", [])
    if recoveries:
        vr_hdr = doc.add_paragraph()
        _bold(vr_hdr, "Voltage Recovery After Reclose", size_pt=10)
        for vr in recoveries:
            if vr.get("error"):
                continue
            shot_num = vr.get("shot_number", "?")
            if vr["recovered"]:
                _body(doc,
                    f"  Shot {shot_num}: Voltage recovered to {vr['post_reclose_pct']:.0f}% "
                    f"of pre-fault level within {vr['recovery_time_ms']:.0f} ms after reclose.  "
                    f"Successful reclosure confirmed.",
                    size_pt=10)
            else:
                _body(doc,
                    f"  Shot {shot_num}: Voltage did NOT recover to "
                    f"{vr['recovery_threshold_pct']:.0f}% within the observation window "
                    f"(post-reclose level: {vr['post_reclose_pct']:.0f}%).  "
                    f"Possible reclose into a still-present fault.",
                    size_pt=10)
        doc.add_paragraph()


def _word_signoff(doc, engineer_name: str = "", engineer_title: str = "",
                  engineer_contact: str = "") -> None:
    """Sign-off block."""
    doc.add_paragraph("Prepared by:")
    doc.add_paragraph()
    p = doc.add_paragraph()
    _bold(p, engineer_name or "[Engineer Name]")
    doc.add_paragraph(engineer_title or "Protection Engineer")
    if engineer_contact:
        doc.add_paragraph(engineer_contact)


# ---------------------------------------------------------------------------
# 3.  Top-level Word report assembler
# ---------------------------------------------------------------------------

def generate_word_report(
    record: EventRecord,
    report: dict,
    outdir: Path,
    stem: str,
    *,
    location: str = "",
    device_type: str = "",
    engineer_name: str = "",
    engineer_title: str = "",
    engineer_contact: str = "",
    waveform_png: Optional[str] = None,
    phasor_png: Optional[str] = None,
    feeder: Optional[dict] = None,
    triage: Optional[dict] = None,
) -> Optional[Path]:
    """
    Generate a Word (.docx) relay event report.

    Parameters
    ----------
    record          EventRecord from the parser
    report          dict from generate_report()
    outdir          output directory Path
    stem            filename stem (e.g. "fault_event")
    waveform_png    path to the saved waveform PNG from plotting.plot_event()
    """
    if not _DOCX_AVAILABLE:
        log.warning(
            "python-docx not installed — cannot generate Word report.\n"
            "Install it with:  pip install python-docx"
        )
        return None

    doc = _DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ---- Title block ----
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _bold(title_p, "Relay Oscillography Event Report", color=_NAVY, size_pt=16)
    doc.add_paragraph()

    # ---- Triage banner ----
    if triage is not None:
        _word_triage(doc, triage)

    # ---- Event information table ----
    _word_event_info_table(doc, record, report, location=location, device_type=device_type)

    # ---- Operations summary table ----
    _word_operations_table(doc, record, report)

    # ---- Peak quantities ----
    _word_peak_quantities(doc, record, report)

    # ---- Waveform figure ----
    _word_waveform(doc, waveform_png)

    # ---- Phasor diagram ----
    _word_phasor(doc, phasor_png)

    # ---- Recloser / feeder section (distribution only) ----
    if feeder is not None:
        _word_reclose_operations(doc, feeder)

    # ---- Digital operations log ----
    _word_digital_log(doc, record, report)

    # ---- Analysis notes ----
    _word_analysis_notes(doc, record, report)

    # ---- Sign-off ----
    _word_signoff(doc, engineer_name, engineer_title, engineer_contact)

    # ---- Save ----
    outdir.mkdir(parents=True, exist_ok=True)
    out_path = outdir / f"{stem}_report.docx"
    doc.save(out_path)
    log.info("Report saved → %s", out_path)
    return out_path
