#!/usr/bin/env python3
"""
wso_impact.py — WSO Reliability Impact Analysis

Processes a folder of normal-day COMTRADE events and quantifies how many
would convert from momentary to sustained outage under EPSS (zero automatic
reclose shots) during Wildfire Safety Operations days.

Classification
--------------
  PERMANENT    Recloser locked out under normal settings.
               WSO makes no difference — already a sustained outage.

  NOT_EXPOSED  Fault cleared without any reclose operation.
               EPSS cannot suppress what didn't happen.

  WSO_EXPOSED  Fault required ≥1 automatic reclose shot to clear.
               Under EPSS (0 shots) this becomes a sustained outage.

Usage
-----
  python3 wso_impact.py ./events/ --devices devices.csv
  python3 wso_impact.py ./events/ --devices devices.csv --response-hours 3
  python3 wso_impact.py ./events/ --no-devices   # no zone grouping
"""

import argparse
import csv
import re
import glob
import json
import os
import sys
from collections import defaultdict
from pathlib import Path
from typing import Optional

_HERE = Path(__file__).parent

from .comtrade_parser import COMTRADEParser
from .analysis import compute_event_summary
from .feeder_analysis import compute_feeder_summary


# ---------------------------------------------------------------------------
# Device registry
# ---------------------------------------------------------------------------

def _normalize(s: str) -> str:
    return s.strip().lower().replace("-", "").replace("_", "").replace(" ", "")


def load_registry(csv_path: str) -> dict:
    """
    Load devices.csv into a dict keyed by normalized device_id.

    Required columns : device_id, zone, risk_tier
    Optional columns : station, feeder, customers_served, aliases

    `aliases` is how real events reach a device. An event's device id is
    whatever was typed into the relay's settings — CFG line 1, field 2 — and it
    is rarely the name the device has on a one-line. List those strings here,
    separated by `;` or `|`, and every one of them resolves to this device.
    Aliases are indexed alongside the id itself, so an event that already
    carries the canonical name still matches.
    """
    registry = {}
    claimed = {}
    with open(csv_path, newline="", encoding="utf-8-sig") as fh:
        reader = csv.DictReader(fh)
        for row in reader:
            row = {k.strip(): v.strip() for k, v in row.items()}
            did = row.get("device_id", "").strip()
            if not did:
                continue
            entry = {
                "device_id":        did,
                "station":          row.get("station", ""),
                "feeder":           row.get("feeder", ""),
                "zone":             row.get("zone", "UNKNOWN"),
                "risk_tier":        int(row.get("risk_tier", 0) or 0),
                "customers_served": int(row.get("customers_served", 0) or 0),
                "aliases":          _split_aliases(row.get("aliases", "")),
            }
            for name in [did, *entry["aliases"]]:
                key = _normalize(name)
                if not key:
                    continue
                # Two devices claiming one alias would silently send events to
                # whichever was read last. First wins, and it is reported.
                if key in registry and claimed.get(key) != did:
                    entry.setdefault("alias_conflicts", []).append(
                        (name, claimed.get(key, "")))
                    continue
                registry[key] = entry
                claimed[key] = did
    return registry


def _split_aliases(raw: str) -> list:
    return [a.strip() for a in re.split(r"[;|]", raw or "") if a.strip()]


def lookup_device(registry: dict, device_id: str, station: str = "") -> Optional[dict]:
    """Case-insensitive match on device_id, then fall back to station name."""
    if key := _normalize(device_id):
        if key in registry:
            return registry[key]
    if station:
        if key2 := _normalize(station):
            if key2 in registry:
                return registry[key2]
    return None


# ---------------------------------------------------------------------------
# Event classification
# ---------------------------------------------------------------------------

PERMANENT      = "PERMANENT"
WSO_EXPOSED    = "WSO_EXPOSED"
NOT_EXPOSED    = "NOT_EXPOSED"
EPSS_CANDIDATE = "EPSS_CANDIDATE"
INDETERMINATE  = "INDETERMINATE"

# Longest reclose dead time worth waiting for before calling a record
# inconclusive. Records that end sooner than this after a trip cannot show
# whether a reclose followed.
_MAX_DEAD_TIME_MS = 5000.0


# key → (order, label, what the record shows, what EPSS does to it, confidence)
_CLASS_TABLE = [
    ("EPSS_CANDIDATE", "Rode through",
     "Fault current recorded, but this device never tripped — normally a "
     "downstream fuse cleared it, so only the lateral went out.",
     "More sensitive EPSS settings may trip this device instead, dropping "
     "everything downstream of it with no reclose. A NEW outage that does not "
     "happen today, and usually a much larger one.",
     "Needs the device's normal and EPSS pickup settings to confirm."),
    ("WSO_EXPOSED", "Momentary → sustained",
     "Tripped, then reclosed successfully — customers saw a blink.",
     "With reclosing disabled the first trip is the last: the blink becomes a "
     "sustained outage until a crew patrols and closes.",
     "Confirmed from the record — a reclose was observed."),
    ("PERMANENT", "Already sustained",
     "Tripped and locked out under normal settings.",
     "Nothing changes. It was already a sustained outage with a truck roll.",
     "Confirmed from the record — lockout observed."),
    ("INDETERMINATE", "Unknown",
     "Tripped, but the record ends before any reclose could have occurred.",
     "Cannot be determined. It is either a momentary that would convert, or a "
     "lockout that would not.",
     "Unresolved. Longer records, or the device's reclose settings, would "
     "settle it."),
    ("NOT_EXPOSED", "No change",
     "No fault current and no device operation.",
     "Nothing for EPSS to act on.",
     "Confirmed from the record."),
]


def class_table() -> list:
    """The EPSS classification scheme as data, for display and export."""
    return [{"key": k, "order": i, "label": lbl, "observed": obs,
             "epss_effect": eff, "confidence": conf}
            for i, (k, lbl, obs, eff, conf) in enumerate(_CLASS_TABLE)]


def classify_event(feeder_data: Optional[dict],
                   summary: Optional[dict] = None,
                   record_end_ms: Optional[float] = None) -> str:
    """
    What this event becomes on an EPSS day.

    EPSS does two things: it disables reclosing, and it makes the relay more
    sensitive and faster. Both matter, and they create different outcomes:

      EPSS_CANDIDATE  Fault current recorded but the device did not trip on
                      normal-day settings — typically a downstream fuse cleared
                      it, so only the lateral went out. Under more sensitive
                      EPSS settings this same fault may trip the recloser,
                      taking out everything downstream of it, with no reclose.
                      A NEW outage that does not exist on a normal day, and
                      usually a much larger one. Confirming it needs the
                      device's normal and EPSS pickup settings.

      WSO_EXPOSED     Tripped and reclosed successfully — a momentary blink
                      today, a sustained outage once reclosing is removed.

      PERMANENT       Locked out on normal settings. Already sustained.

      INDETERMINATE   Tripped, but the record ends before any reclose could
                      have occurred. Cannot tell which of the two above it is.

      NOT_EXPOSED     No fault current and no operation. Nothing to convert.
                      A balanced load step with the voltage intact (fault_type
                      LOAD) counts as no fault current — a tie closing onto a
                      restored section is not something EPSS converts.
    """
    if feeder_data is None:
        return NOT_EXPOSED
    seq = feeder_data.get("reclose_sequence")
    if seq is None:
        return NOT_EXPOSED

    if seq.locked_out:
        return PERMANENT

    if any(s.outcome == "RECLOSED" for s in seq.shots):
        return WSO_EXPOSED

    # No trip at all. If fault current was present, the device rode through
    # something — the case this whole analysis exists to find.
    if seq.total_shots == 0:
        if (summary and summary.get("fault_inception_s") is not None
                and summary.get("fault_type") != "LOAD"):
            return EPSS_CANDIDATE
        return NOT_EXPOSED

    # Tripped once with no reclose in the record. Whether that is a lockout or
    # a reclose we simply did not record depends on how long the record ran.
    if summary and record_end_ms is not None:
        trip_s = summary.get("trip_time_s")
        if trip_s is not None and (record_end_ms - trip_s * 1000.0) < _MAX_DEAD_TIME_MS:
            return INDETERMINATE
    return NOT_EXPOSED


# ---------------------------------------------------------------------------
# Folder analysis
# ---------------------------------------------------------------------------

def _find_cfg(folder: str) -> list:
    found = []
    for pat in ("*.cfg", "*.CFG", "*.cff", "*.CFF"):
        found.extend(glob.glob(os.path.join(folder, pat)))
        found.extend(glob.glob(os.path.join(folder, "**", pat), recursive=True))
    return sorted(set(found))


def analyze_folder(
    folder: str,
    registry: dict,
    response_hours: float = 2.0,
    epss_tiers: list = None,
    verbose: bool = True,
) -> dict:
    if epss_tiers is None:
        epss_tiers = [2, 3]

    files = _find_cfg(folder)
    if not files:
        print(f"No COMTRADE files found under: {folder}", file=sys.stderr)
        return {}

    parser = COMTRADEParser()

    zone_buckets  = defaultdict(lambda: {"devices_seen": set(), "customers_served": 0,
                                          "risk_tiers": set(), "events": []})
    unmatched     = []
    parse_errors  = []

    print(f"Processing {len(files)} file(s) ...")

    for filepath in files:
        try:
            record = parser.parse(filepath)
        except Exception as exc:
            parse_errors.append({"file": filepath, "error": str(exc)})
            if verbose:
                print(f"  ERROR   {os.path.basename(filepath)}: {exc}")
            continue

        summary     = compute_event_summary(record)
        feeder_data = compute_feeder_summary(record)
        cls         = classify_event(feeder_data)

        seq = feeder_data.get("reclose_sequence") if feeder_data else None
        event = {
            "file":           os.path.basename(filepath),
            "device_id":      record.metadata.get("rec_dev_id", ""),
            "station":        record.metadata.get("station_name", ""),
            "fault_type":     summary.get("fault_type", "UNKNOWN"),
            "trip_delay_ms":  summary.get("trip_delay_ms"),
            "total_shots":    (seq.total_shots if seq else 0),
            "classification": cls,
        }

        device = lookup_device(registry, record.metadata.get("rec_dev_id", ""), record.metadata.get("station_name", ""))
        if device is None:
            unmatched.append(event)
            if verbose:
                lbl = {"PERMANENT": "PERM", "WSO_EXPOSED": "WSO!", "NOT_EXPOSED": "ok  "}[cls]
                print(f"  [{lbl}] UNMATCHED  {record.metadata.get("rec_dev_id", "")!r}")
            continue

        zone = device["zone"]
        zb   = zone_buckets[zone]
        if device["device_id"] not in zb["devices_seen"]:
            zb["devices_seen"].add(device["device_id"])
            zb["customers_served"] += device["customers_served"]
        zb["risk_tiers"].add(device["risk_tier"])
        event["device_entry"] = device
        zb["events"].append(event)

        if verbose:
            lbl = {"PERMANENT": "PERM", "WSO_EXPOSED": "WSO!", "NOT_EXPOSED": "ok  "}[cls]
            print(f"  [{lbl}] {record.metadata.get("rec_dev_id", "") or record.metadata.get("station_name", ""):<20s}  "
                  f"{event['fault_type']:6s}  zone={zone}  shots={event['total_shots']}")

    # ── Aggregate ────────────────────────────────────────────────────────────

    def _tally(events):
        counts = {PERMANENT: 0, WSO_EXPOSED: 0, NOT_EXPOSED: 0}
        ft = defaultdict(int)
        for e in events:
            counts[e["classification"]] += 1
            ft[e["fault_type"]] += 1
        return counts, dict(ft)

    zones_out = {}
    sys_perm = sys_exposed = sys_not = 0
    sys_customers = sys_cust_hours = 0

    for zone_name, zb in sorted(zone_buckets.items()):
        counts, ft = _tally(zb["events"])
        tiers        = sorted(zb["risk_tiers"])
        epss_active  = any(t in epss_tiers for t in tiers)
        exposed      = counts[WSO_EXPOSED] if epss_active else 0
        cust         = zb["customers_served"]
        cust_hours   = exposed * cust * response_hours

        zones_out[zone_name] = {
            "risk_tiers":                     tiers,
            "epss_active":                    epss_active,
            "device_count":                   len(zb["devices_seen"]),
            "customers_served":               cust,
            "events_total":                   len(zb["events"]),
            "permanent":                      counts[PERMANENT],
            "wso_exposed":                    exposed,
            "not_exposed":                    counts[NOT_EXPOSED],
            "fault_type_breakdown":           ft,
            "est_customer_hours_per_wso_day": cust_hours,
        }
        sys_perm     += counts[PERMANENT]
        sys_exposed  += exposed
        sys_not      += counts[NOT_EXPOSED]
        sys_customers += cust
        sys_cust_hours += cust_hours

    um_counts, um_ft = _tally(unmatched)
    total = sys_perm + sys_exposed + sys_not + len(unmatched)

    return {
        "folder":         folder,
        "files_found":    len(files),
        "parse_errors":   parse_errors,
        "epss_max_shots": 0,
        "epss_tiers":     epss_tiers,
        "response_hours": response_hours,
        "zones":          zones_out,
        "unmatched": {
            "events":      len(unmatched),
            "permanent":   um_counts[PERMANENT],
            "wso_exposed": um_counts[WSO_EXPOSED],
            "not_exposed": um_counts[NOT_EXPOSED],
            "fault_types": um_ft,
            "device_ids":  sorted({e["device_id"] for e in unmatched}),
        },
        "system": {
            "events_total":                   total,
            "events_matched":                 total - len(unmatched),
            "events_unmatched":               len(unmatched),
            "permanent":                      sys_perm,
            "wso_exposed":                    sys_exposed,
            "not_exposed":                    sys_not,
            "wso_exposed_pct":                round(100 * sys_exposed / total, 1) if total else 0.0,
            "customers_served":               sys_customers,
            "est_customer_hours_per_wso_day": sys_cust_hours,
        },
    }


# ---------------------------------------------------------------------------
# Console summary
# ---------------------------------------------------------------------------

def print_summary(results: dict):
    s = results["system"]
    W = 66
    print()
    print("=" * W)
    print("  WSO RELIABILITY IMPACT ANALYSIS")
    print("=" * W)
    print(f"  Folder       : {results['folder']}")
    print(f"  Files        : {results['files_found']}")
    if results["parse_errors"]:
        print(f"  Parse errors : {len(results['parse_errors'])}")
    print(f"  EPSS setting : 0 shots  (tiers {results['epss_tiers']})")
    print(f"  Avg response : {results['response_hours']:.1f} h")

    for zone_name, z in results["zones"].items():
        tiers = "/".join(str(t) for t in z["risk_tiers"])
        epss  = "EPSS active" if z["epss_active"] else "EPSS inactive"
        print()
        print(f"  ZONE {zone_name}  (Tier {tiers}, {epss})")
        print(f"    Devices  : {z['device_count']}   Customers: {z['customers_served']:,}")
        print(f"    Events   : {z['events_total']}")
        print(f"      Already permanent  : {z['permanent']}")
        print(f"      No reclose needed  : {z['not_exposed']}")
        print(f"      WSO-exposed        : {z['wso_exposed']}  ← sustained outage under EPSS")
        if z["wso_exposed"] and z["customers_served"]:
            print(f"      Est. cust-hrs/day  : {z['est_customer_hours_per_wso_day']:,.0f}")
        if z["fault_type_breakdown"]:
            ft = "  ".join(f"{k}:{v}" for k, v in sorted(z["fault_type_breakdown"].items()))
            print(f"      Fault types        : {ft}")

    um = results["unmatched"]
    if um["events"]:
        print()
        print(f"  UNREGISTERED ({um['events']} events — excluded from zone totals)")
        print(f"    WSO-exposed : {um['wso_exposed']}")
        ids = um["device_ids"]
        print(f"    Device IDs  : {', '.join(ids[:6])}" + ("..." if len(ids) > 6 else ""))

    print()
    print("─" * W)
    print("  SYSTEM TOTAL (registered devices only)")
    print(f"    Events analyzed       : {s['events_total']}")
    print(f"    Matched to registry   : {s['events_matched']}")
    print(f"    Already permanent     : {s['permanent']}")
    print(f"    No reclose needed     : {s['not_exposed']}")
    print(f"    WSO-exposed           : {s['wso_exposed']}  ({s['wso_exposed_pct']}%)")
    if s["customers_served"]:
        print(f"    Customers in analysis : {s['customers_served']:,}")
    if s["est_customer_hours_per_wso_day"]:
        print(f"    Est. cust-hrs / day   : {s['est_customer_hours_per_wso_day']:,.0f}")
    print("=" * W)


# ---------------------------------------------------------------------------
# JSON export
# ---------------------------------------------------------------------------

def write_json(results: dict, path: str):
    def _serial(v):
        if isinstance(v, set):
            return sorted(v)
        return str(v)
    with open(path, "w", encoding="utf-8") as fh:
        json.dump(results, fh, indent=2, default=_serial)
    print(f"  JSON → {path}")


# ---------------------------------------------------------------------------
# Word report
# ---------------------------------------------------------------------------

def write_word_report(results: dict, path: str):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("  python-docx not installed — skipping Word report")
        return

    def _rgb(r, g, b):   return RGBColor(r, g, b)
    _NAVY  = _rgb(0x1A, 0x3A, 0x6B)
    _RED   = _rgb(0xC0, 0x39, 0x2B)
    _WHITE = _rgb(0xFF, 0xFF, 0xFF)

    def _shade(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def _widths(table, cms):
        for row in table.rows:
            for cell, w in zip(row.cells, cms):
                cell.width = Cm(w)

    def _bold(para, text, color=None, size=11):
        r = para.add_run(text); r.bold = True; r.font.size = Pt(size)
        if color: r.font.color.rgb = color
        return r

    def _text(para, text, color=None, size=10):
        r = para.add_run(text); r.font.size = Pt(size)
        if color: r.font.color.rgb = color
        return r

    def _heading(doc, title):
        p = doc.add_paragraph()
        _bold(p, title, color=_NAVY, size=11)

    def _kv_table(doc, rows_data, widths=(8.0, 8.0)):
        tbl = doc.add_table(rows=len(rows_data), cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        _widths(tbl, list(widths))
        for i, (label, value, highlight) in enumerate(rows_data):
            row = tbl.rows[i]
            _shade(row.cells[0], "F2F2F2")
            lp = row.cells[0].paragraphs[0]
            lr = lp.add_run(label); lr.bold = True; lr.font.size = Pt(10)
            vp = row.cells[1].paragraphs[0]
            vr = vp.add_run(str(value)); vr.font.size = Pt(10)
            if highlight: vr.bold = True; vr.font.color.rgb = _RED
        return tbl

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.0)
        sec.left_margin = sec.right_margin = Cm(2.5)

    # Title
    p = doc.add_paragraph()
    _bold(p, "WSO Reliability Impact Analysis", color=_NAVY, size=16)
    p2 = doc.add_paragraph()
    s  = results["system"]
    _text(p2,
          f"Folder: {results['folder']}  |  "
          f"Files: {results['files_found']}  |  "
          f"EPSS: 0 shots (Tiers {results['epss_tiers']})  |  "
          f"Avg response: {results['response_hours']:.1f} h",
          size=9)
    doc.add_paragraph()

    # System summary
    _heading(doc, "System Summary")
    _kv_table(doc, [
        ("Events analyzed",                   s["events_total"],                            False),
        ("Matched to device registry",         s["events_matched"],                          False),
        ("Unmatched (excluded from totals)",   s["events_unmatched"],                        False),
        ("Already permanent — no WSO effect",  s["permanent"],                               False),
        ("No reclose needed — no WSO effect",  s["not_exposed"],                             False),
        ("WSO-exposed (becomes sustained)",
         f"{s['wso_exposed']}  ({s['wso_exposed_pct']}%)",                                   True),
        ("Customers covered by registry",      f"{s['customers_served']:,}",                 False),
        ("Est. customer-hours / WSO day",
         f"{s['est_customer_hours_per_wso_day']:,.0f}" if s["est_customer_hours_per_wso_day"] else "—",
         False),
    ])
    doc.add_paragraph()

    # Per-zone
    for zone_name, z in results["zones"].items():
        tiers = "/".join(str(t) for t in z["risk_tiers"])
        epss  = "EPSS ACTIVE" if z["epss_active"] else "EPSS inactive"
        _heading(doc, f"Zone: {zone_name}  —  Tier {tiers}  —  {epss}")

        _kv_table(doc, [
            ("Devices in zone",       z["device_count"],                     False),
            ("Customers served",      f"{z['customers_served']:,}",          False),
            ("Events analyzed",       z["events_total"],                     False),
            ("Already permanent",     z["permanent"],                        False),
            ("No reclose needed",     z["not_exposed"],                      False),
            ("WSO-exposed",           z["wso_exposed"],                      z["wso_exposed"] > 0),
            ("Est. cust-hrs / WSO day",
             f"{z['est_customer_hours_per_wso_day']:,.0f}"
             if z["est_customer_hours_per_wso_day"] else "—",                False),
        ])

        if z["fault_type_breakdown"]:
            p3 = doc.add_paragraph()
            ft = "  ".join(f"{k}: {v}" for k, v in sorted(z["fault_type_breakdown"].items()))
            _text(p3, f"Fault type breakdown: {ft}", size=9)
        doc.add_paragraph()

    # Unregistered
    um = results["unmatched"]
    if um["events"]:
        _heading(doc, "Unregistered Devices")
        p = doc.add_paragraph()
        _text(p,
              f"{um['events']} event(s) could not be matched to the device registry "
              f"and are excluded from zone totals.  WSO-exposed in this group: {um['wso_exposed']}.  "
              f"Add these device IDs to devices.csv to include them in future runs.",
              size=10)
        if um["device_ids"]:
            p2 = doc.add_paragraph()
            _text(p2, "Unrecognized IDs: " + ", ".join(um["device_ids"]), size=9)
        doc.add_paragraph()

    # Methodology
    _heading(doc, "Methodology")
    p = doc.add_paragraph()
    _text(p,
          "Events are classified as WSO-exposed if the fault required one or more automatic "
          "reclose operations to clear under normal operating settings.  Under EPSS (zero "
          "automatic reclose shots), these events would result in sustained customer outages "
          "requiring manual field crew restoration.  Events that locked out under normal "
          "settings are already sustained outages and are unaffected by EPSS.  Events that "
          "cleared on the initial trip without a reclose are similarly unaffected.  "
          f"Customer-hour estimates assume an average field crew response time of "
          f"{results['response_hours']:.1f} hours per event.",
          size=9)

    doc.save(path)
    print(f"  DOCX → {path}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    p = argparse.ArgumentParser(
        prog="wso-impact",
        description="WSO Reliability Impact Analysis — quantify EPSS effect on distribution reliability",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
examples:
  python3 wso_impact.py ./events/ --devices devices.csv
  python3 wso_impact.py ./events/ --devices devices.csv --response-hours 3
  python3 wso_impact.py ./events/ --no-devices
        """,
    )
    p.add_argument("folder",
                   help="Folder containing COMTRADE event files (searched recursively)")
    p.add_argument("--devices", metavar="CSV",
                   help="Path to device registry CSV")
    p.add_argument("--no-devices", action="store_true",
                   help="Run without a device registry — skips zone grouping, still classifies events")
    p.add_argument("--response-hours", type=float, default=2.0, metavar="HRS",
                   help="Average crew response time in hours for customer-hour estimates (default 2.0)")
    p.add_argument("--epss-tiers", type=int, nargs="+", default=[2, 3], metavar="N",
                   help="Risk tiers that receive EPSS / zero-shot treatment (default: 2 3)")
    p.add_argument("--out-dir", metavar="DIR",
                   help="Output directory (default: <folder>/wso_output/)")
    p.add_argument("--quiet", action="store_true",
                   help="Suppress per-file progress output")

    args = p.parse_args()

    if not os.path.isdir(args.folder):
        print(f"Error: folder not found — {args.folder}", file=sys.stderr)
        sys.exit(1)

    registry = {}
    if args.devices:
        if not os.path.isfile(args.devices):
            print(f"Error: devices file not found — {args.devices}", file=sys.stderr)
            sys.exit(1)
        registry = load_registry(args.devices)
        print(f"Registry: {len(registry)} device(s) from {args.devices}")
    elif not args.no_devices:
        print("Note: no --devices file given. Pass --no-devices to suppress this message.",
              file=sys.stderr)

    results = analyze_folder(
        args.folder,
        registry,
        response_hours=args.response_hours,
        epss_tiers=args.epss_tiers,
        verbose=not args.quiet,
    )
    if not results:
        sys.exit(1)

    print_summary(results)

    out_dir = args.out_dir or os.path.join(args.folder, "wso_output")
    os.makedirs(out_dir, exist_ok=True)
    write_json(results,        os.path.join(out_dir, "wso_impact.json"))
    write_word_report(results, os.path.join(out_dir, "wso_impact_report.docx"))
    print(f"\nOutputs → {out_dir}")


if __name__ == "__main__":
    main()
